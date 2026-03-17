# exportar_docx.py
import json
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


# ── Colores del Formato Maestro Lord Byron ────────────────────────────────────
VERDE_PRINCIPAL  = RGBColor(0x16, 0x6C, 0x52)   # #166C52
VERDE_SECUNDARIO = RGBColor(0x3F, 0x7A, 0x63)   # #3f7a63
ROSA_OBJECTIVE   = RGBColor(0xFF, 0x33, 0x66)   # rgb(255,51,102)
AZUL_INSTRUC     = RGBColor(0x33, 0x66, 0xFF)   # #3366FF
BLANCO           = RGBColor(0xFF, 0xFF, 0xFF)


def _set_cell_bg(cell, color_hex: str):
    """Pone color de fondo a una celda de tabla."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  color_hex)
    tcPr.append(shd)


def _add_label(doc: Document, texto: str, color_hex: str, font_size: int = 14):
    """
    Agrega una etiqueta estilo 'badge' usando una tabla de 1 celda
    con fondo de color (simula las etiquetas verdes del HTML).
    """
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell  = table.cell(0, 0)
    _set_cell_bg(cell, color_hex.lstrip("#"))
    p     = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run   = p.add_run(texto)
    run.bold       = True
    run.font.color.rgb = BLANCO
    run.font.size  = Pt(font_size)
    run.font.name  = "Arial"
    doc.add_paragraph()  # espacio después


def _add_session_header(doc: Document, session_num: str,
                        reading_title: str, chapter: str):
    """SESSION + título + subtítulo centrados."""
    # SESSION badge — centrado
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"SESSION {session_num}")
    run.bold           = True
    run.font.size      = Pt(22)
    run.font.color.rgb = VERDE_PRINCIPAL
    run.font.name      = "Montserrat"

    # Título de la lectura
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(reading_title.upper())
    r2.bold           = True
    r2.font.size      = Pt(18)
    r2.font.color.rgb = VERDE_PRINCIPAL
    r2.font.name      = "Arial"

    # Chapter / subtítulo
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run(chapter)
    r3.font.size      = Pt(13)
    r3.font.color.rgb = VERDE_PRINCIPAL
    r3.font.name      = "Arial"
    doc.add_paragraph()


def _add_objective(doc: Document, text: str):
    """OBJECTIVE en rosa + texto en Arial."""
    # Header OBJECTIVE
    p = doc.add_paragraph()
    run = p.add_run("OBJECTIVE")
    run.bold           = True
    run.font.size      = Pt(12)
    run.font.color.rgb = ROSA_OBJECTIVE
    run.font.name      = "Arial"

    # Texto del objetivo
    if not text.strip().startswith("Students will"):
        text = "Students will " + text.lstrip()
    p2 = doc.add_paragraph(text)
    p2.runs[0].font.size = Pt(11)
    p2.runs[0].font.name = "Arial"
    doc.add_paragraph()


def _add_activity(doc: Document, act: dict):
    """ACTIVITY badge + instrucciones + preguntas."""
    number = act.get("number", "01")
    emoji  = act.get("emoji",  "")
    label  = act.get("label",  f"ACTIVITY {number}")
    if label.startswith("ACTIVITY"):
        label = f"ACTIVITY {number}"

    # Badge de actividad (verde secundario)
    _add_label(doc, f"{label} {emoji}", "#3f7a63", font_size=13)

    # Instructions header
    instructions = act.get("instructions", [])
    if instructions:
        p = doc.add_paragraph()
        run = p.add_run("Instructions:")
        run.bold           = True
        run.font.size      = Pt(11)
        run.font.color.rgb = AZUL_INSTRUC
        run.font.name      = "Arial"

        for inst in instructions:
            words = inst.split(" ", 1)
            bullet = doc.add_paragraph(style="List Bullet")
            if len(words) == 2:
                r1 = bullet.add_run(words[0] + " ")
                r1.bold      = True
                r1.font.name = "Arial"
                r1.font.size = Pt(11)
                r2 = bullet.add_run(words[1])
                r2.font.name = "Arial"
                r2.font.size = Pt(11)
            else:
                r = bullet.add_run(inst)
                r.bold      = True
                r.font.name = "Arial"
                r.font.size = Pt(11)

    # Preguntas
    questions = act.get("questions", [])
    if questions:
        for q in questions:
            bullet = doc.add_paragraph(style="List Bullet")
            run = bullet.add_run(q)
            run.font.name = "Arial"
            run.font.size = Pt(11)

    doc.add_paragraph()


def exportar_docx(approved_content: str, session_num: str,
                  reading_title: str, chapter_topic: str,
                  output_path: str):
    """
    Genera el archivo .docx desde el JSON aprobado.
    Respeta el Formato Maestro Lord Byron.
    """
    try:
        data = json.loads(approved_content)
    except (json.JSONDecodeError, TypeError):
        print("  ⚠️  No se pudo exportar a DOCX: contenido no es JSON válido.")
        return

    doc = Document()

    # Márgenes de página (2 cm todos los lados)
    for section in doc.sections:
        section.top_margin    = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin   = Inches(1.0)
        section.right_margin  = Inches(1.0)

    # ── Estructura obligatoria del Formato Maestro ────────────────────────────

    # 1–3. SESSION + Título + Chapter
    _add_session_header(doc, session_num, reading_title, chapter_topic)

    # 4. LEADING FOR LEARNING
    _add_label(doc, "LEADING FOR LEARNING", "#166C52", font_size=14)

    # 5. OBJECTIVE (Leading)
    _add_objective(doc, data.get("leading_objective", "Students will activate prior knowledge."))

    # 6. ACTIVITY (Leading)
    if "leading_activity" in data:
        _add_activity(doc, data["leading_activity"])

    # 7. BUILDING UP KNOWLEDGE
    _add_label(doc, "BUILDING UP KNOWLEDGE", "#166C52", font_size=14)

    # 8. OBJECTIVE (Building)
    _add_objective(doc, data.get("building_objective", "Students will identify key concepts."))

    # 9–10. BEFORE READING
    _add_label(doc, "BEFORE READING", "#166C52", font_size=14)
    if "before_reading" in data:
        _add_activity(doc, data["before_reading"])

    # 11–12. DURING READING
    _add_label(doc, "DURING READING", "#166C52", font_size=14)
    if "during_reading" in data:
        _add_activity(doc, data["during_reading"])

    # 13–14. AFTER READING
    _add_label(doc, "AFTER READING", "#166C52", font_size=14)
    if "after_reading" in data:
        _add_activity(doc, data["after_reading"])

    # Extra opcional
    if "extra_activity" in data:
        _add_activity(doc, data["extra_activity"])

    # ── Guardar ───────────────────────────────────────────────────────────────
    doc.save(output_path)
    print(f"  [✔] DOCX guardado: {output_path}")